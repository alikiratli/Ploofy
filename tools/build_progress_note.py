"""content/ilerleme-notu.md -> "ilerleme notu.docx"

İlerleme notunun kaynağı markdown; docx üretilen çıktı. Sebebi: metin
okunabilir ve git'te satır satır izlenebilir kalsın, ama açılan dosya
Word belgesi olsun.

Kullanım:
    python tools/build_progress_note.py

Desteklenen biçim: # / ## / ### başlıklar, "- " madde işaretleri,
"1. " numaralı maddeler, **kalın** ve `kod` satır içi işaretleri.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INK = RGBColor(0x3A, 0x2A, 0x1E)
MUTED = RGBColor(0x7A, 0x6A, 0x5C)

# **kalın** ve `kod` parçalarını ayırır; yakalama grupları sayesinde
# ayıraçlar da listede kalıyor.
INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_runs(paragraph, text):
    """Satır içi kalın/kod işaretlerini gerçek biçimlendirmeye çevirir."""
    for part in INLINE.split(text):
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)

        run.font.color.rgb = INK


def build(source: Path, target: Path) -> None:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)

    lines = source.read_text(encoding="utf-8").splitlines()

    # Markdown'da satır kaydırma sırf okunurluk için; Word'de her kaynak
    # satırı ayrı paragraf olursa cümleler ortasından bölünüyor. Bu yüzden
    # ardışık satırlar boş satıra ya da yeni bir bloğa kadar biriktiriliyor.
    buffer: list[str] = []
    style = ""

    # Ana başlık ile ilk bölüm başlığı arasındaki satırlar künye (tarih,
    # depo adresi): gövde metninden ayrılsın diye gri ve küçük.
    in_masthead = False

    def flush():
        nonlocal buffer, style
        if not buffer:
            return

        text = " ".join(buffer)
        if style == "MASTHEAD":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
        else:
            add_runs(document.add_paragraph(style=style), text)

        buffer = []
        style = ""

    def start(new_style, text):
        nonlocal buffer, style
        flush()
        style = new_style
        buffer = [text]

    for line in lines:
        stripped = line.strip()

        if not stripped:
            flush()
            continue

        if stripped.startswith("### "):
            flush()
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            flush()
            in_masthead = False
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            flush()
            heading = document.add_heading(stripped[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            in_masthead = True
        elif stripped.startswith("- "):
            start("List Bullet", stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            start("List Number", re.sub(r"^\d+\.\s", "", stripped))
        elif in_masthead:
            # Künyenin her satırı kendi başına duruyor.
            start("MASTHEAD", stripped)
            flush()
        elif buffer:
            # Açık bir bloğun devam satırı — madde de olabilir, paragraf da.
            buffer.append(stripped)
        else:
            start("Normal", stripped)

    flush()

    document.save(target)
    print(f"{target.name} yazildi ({len(lines)} satir kaynaktan)")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    source = root / "content" / "ilerleme-notu.md"
    target = root / "ilerleme notu.docx"

    if not source.exists():
        sys.exit(f"Kaynak bulunamadi: {source}")

    build(source, target)


if __name__ == "__main__":
    main()
